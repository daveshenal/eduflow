from typing import List, Dict, Any, Optional
import uuid
from datetime import datetime, timezone
from io import BytesIO
import logging
from dataclasses import dataclass
from enum import Enum
import PyPDF2

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from config.settings import settings
from langchain.schema import Document

# Azure services
from azure.core.exceptions import ResourceNotFoundError
from app.adapters.azure_search import get_search_client, get_search_index_client
from app.adapters.azure_blob import get_blob_service_client

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Index/container naming (matches AIIndex in ai_index.py)
INDEX_PREFIX = "ai-index-"
CONTAINER_PREFIX = "ai-"


def _index_name(index_id: str) -> str:
    """Get search index name for index_id."""
    return f"{INDEX_PREFIX}{index_id}"


def _container_name(index_id: str) -> str:
    """Get blob container name for index_id."""
    return f"{CONTAINER_PREFIX}{index_id}"


class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    UNKNOWN = "unknown"


@dataclass
class ProcessingResult:
    """Result of document processing"""
    success: bool
    message: str
    chunks_created: int = 0
    chunks_failed: int = 0
    total_chunks: int = 0
    error: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class DocumentProcessor:
    """Main document processing class using LangChain."""

    def __init__(
        self,
        index_id: str,
        embedding_deployment: str = None,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        ):
        
        self.index_id = index_id
        self.search_client = get_search_client(_index_name(index_id))

        deployment_name = embedding_deployment or settings.AZURE_OPENAI_EMBEDDING_DEPLOYMENT
        self.embeddings = AzureOpenAIEmbeddings(
            deployment=deployment_name,
            model=deployment_name,
            azure_endpoint=settings.AZURE_OPENAI_ENDPOINT,
            api_key=settings.AZURE_OPENAI_KEY,
            openai_api_version=settings.AZURE_OPENAI_API_VERSION,
            chunk_size=200,
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def _get_file_type(self, filename: str) -> FileType:
        """Determine file type from filename"""
        extension = filename.lower().split(".")[-1]
        try:
            return FileType(extension)
        except ValueError:
            return FileType.UNKNOWN

    def _load_document_from_bytes(self, filename: str, data: bytes) -> List[Document]:
        """Load document from bytes using appropriate loader"""
        file_type = self._get_file_type(filename)

        try:
            if file_type == FileType.PDF:
                return self._load_pdf_from_bytes(data, filename)
            elif file_type in [FileType.DOCX, FileType.DOC]:
                return self._load_docx_from_bytes(data, filename)
            elif file_type == FileType.TXT:
                return self._load_txt_from_bytes(data, filename)
            else:
                return self._load_txt_from_bytes(data, filename)
        except Exception as e:
            logger.error(f"Error loading {filename}: {e}")
            return []

    def _load_pdf_from_bytes(self, data: bytes, filename: str) -> List[Document]:
        """Load PDF from bytes"""
        try:
            pdf_file = BytesIO(data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            documents = []
            for _, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    doc = Document(
                        page_content=page_text,
                        metadata={
                            "doc_id": str(uuid.uuid4()),
                            "filename": filename,
                            "file_type": FileType.PDF.value,
                        },
                    )
                    documents.append(doc)

            return documents
        except Exception as e:
            error_msg = f"Error extracting PDF text from {filename}: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)

    def _load_docx_from_bytes(self, data: bytes, filename: str) -> List[Document]:
        """Load DOCX from bytes"""
        try:
            import docx
            doc_file = BytesIO(data)
            doc = docx.Document(doc_file)

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            return [
                Document(
                    page_content=text,
                    metadata={
                        "doc_id": str(uuid.uuid4()),
                        "filename": filename,
                        "file_type": FileType.DOCX.value,
                    },
                )
            ]
        except ImportError as e:
            error_msg = f"Missing required dependency for DOCX processing: {str(e)}. Please install python-docx: pip install python-docx"
            logger.error(error_msg)
            raise Exception(error_msg)
        except Exception as e:
            error_msg = f"Error extracting DOCX text from {filename}: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)

    def _load_txt_from_bytes(self, data: bytes, filename: str) -> List[Document]:
        """Load text from bytes"""
        try:
            text = data.decode("utf-8")
            return [
                Document(
                    page_content=text,
                    metadata={
                        "doc_id": str(uuid.uuid4()),
                        "filename": filename,
                        "file_type": FileType.TXT.value,
                    },
                )
            ]
        except Exception as e:
            logger.error(f"Error decoding text: {e}")
            return []

    def _prepare_documents_for_search(
        self, documents: List[Document], filename: str
     ) -> List[Dict[str, Any]]:
        
        """Prepare documents for Azure Search index (AIIndex schema)."""
        search_docs = []

        texts = [doc.page_content for doc in documents]

        try:
            embeddings = self.embeddings.embed_documents(texts)

            for doc, embedding in zip(documents, embeddings):
                search_doc = {
                    "chunk_id": str(uuid.uuid4()),
                    "parent_id": doc.metadata.get("doc_id"),
                    "category": None,
                    "source_name": filename,
                    "source_path": filename,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "title": filename,
                    "content": doc.page_content,
                    "content_vector": embedding,
                }
                search_docs.append(search_doc)

        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return []

        return search_docs

    def process_document(
        self, filename: str, data: bytes
        ) -> ProcessingResult:
        
        """Process a document through the complete pipeline."""
        logger.info(f"Processing document: {filename} for index_id: {self.index_id}")

        try:
            documents = self._load_document_from_bytes(filename, data)
            if not documents:
                return ProcessingResult(
                    success=False,
                    message="No content could be extracted from the file",
                    error="Failed to load document",
                )

            logger.info(f"Loaded {len(documents)} document(s)")

            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Created {len(chunks)} chunks")

            if not chunks:
                return ProcessingResult(
                    success=False,
                    message="No chunks could be created from the document",
                    error="Failed to create chunks",
                )

            search_docs = self._prepare_documents_for_search(chunks, filename)

            if not search_docs:
                return ProcessingResult(
                    success=False,
                    message="Failed to generate embeddings for chunks",
                    error="Embedding generation failed",
                    total_chunks=len(chunks),
                )

            try:
                self.search_client.upload_documents(search_docs)
                logger.info("Successfully uploaded to search index")

                return ProcessingResult(
                    success=True,
                    message=f"Successfully processed {filename}",
                    chunks_created=len(search_docs),
                    chunks_failed=len(chunks) - len(search_docs),
                    total_chunks=len(chunks),
                    metadata=documents[0].metadata if documents else {},
                )

            except Exception as e:
                logger.error(f"Failed to upload to search index: {e}")
                return ProcessingResult(
                    success=False,
                    message="Failed to upload to search index",
                    error=str(e),
                    total_chunks=len(chunks),
                )

        except Exception as e:
            logger.error(f"Error processing document: {e}")
            return ProcessingResult(
                success=False,
                message="Failed to process document",
                error=str(e),
            )


class DocumentManager:
    """High-level document management class."""

    def __init__(self, index_id: str):
        self.index_id = index_id
        self.processor = DocumentProcessor(index_id=index_id)
        self.blob_service_client = get_blob_service_client()
        self.container_name = _container_name(index_id)

    def _validate_index_and_container_exist(self) -> None:
        """Raise error if index or blob container does not exist. Do not create them."""
        index_name = _index_name(self.index_id)
        try:
            get_search_index_client().get_index(index_name)
        except ResourceNotFoundError:
            raise ValueError(
                f"Search index '{index_name}' does not exist. "
            )
        try:
            container_client = self.blob_service_client.get_container_client(
                self.container_name
            )
            container_client.get_container_properties()
        except ResourceNotFoundError:
            raise ValueError(
                f"Blob container '{self.container_name}' does not exist. "
            )

    def upload_and_process(
        self, filename: str, data: bytes
        ) -> Dict[str, Any]:
        
        """Upload file to blob storage and process into search index."""
        self._validate_index_and_container_exist()

        blob_name = filename
        container_client = self.blob_service_client.get_container_client(
            self.container_name
        )
        blob_client = container_client.get_blob_client(blob_name)

        try:
            blob_client.upload_blob(data, overwrite=True)
            logger.info(f"Uploaded original document to blob: {blob_name}")

            processing_result = self.processor.process_document(filename, data)

            return {
                "blob_name": blob_name,
                "blob_url": blob_client.url,
                "processing_result": {
                    "success": processing_result.success,
                    "message": processing_result.message,
                    "chunks_created": processing_result.chunks_created,
                    "chunks_failed": processing_result.chunks_failed,
                    "total_chunks": processing_result.total_chunks,
                    "error": processing_result.error,
                },
                "success": processing_result.success,
            }

        except Exception as e:
            logger.error(f"Error uploading and processing: {e}")
            return {
                "blob_name": blob_name,
                "processing_result": {"error": str(e)},
                "success": False,
            }

    def batch_process(
        self, files: List[tuple]
        ) -> Dict[str, Any]:
        
        """Process multiple files in batch."""
        results = []
        total_success = 0
        total_failed = 0

        for filename, data in files:
            try:
                result = self.upload_and_process(filename, data)
                results.append({"filename": filename, "result": result})

                if result["success"]:
                    total_success += 1
                else:
                    total_failed += 1

            except Exception as e:
                results.append(
                    {
                        "filename": filename,
                        "result": {"error": str(e), "success": False},
                    }
                )
                total_failed += 1

        return {
            "total_processed": len(files),
            "successful": total_success,
            "failed": total_failed,
            "results": results,
        }

    def delete_index_documents(self) -> Dict[str, Any]:
        """Delete all documents in the search index for this index_id."""
        try:
            results = self.processor.search_client.search(
                search_text="*",
                select=["chunk_id"],
                include_total_count=True,
            )

            doc_ids = [{"chunk_id": r["chunk_id"]} for r in results]
            if doc_ids:
                self.processor.search_client.delete_documents(doc_ids)
                logger.info(f"Deleted {len(doc_ids)} documents for index {self.index_id}")
                return {"deleted_count": len(doc_ids), "success": True}
            else:
                return {"deleted_count": 0, "success": True, "message": "No documents found"}

        except Exception as e:
            logger.error(f"Error deleting index documents: {e}")
            return {"error": str(e), "success": False}
